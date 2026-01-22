import io
from pathlib import Path
from typing import Union

# Try to import PDF libraries
try:
    import pdfplumber
    PDF_LIBRARY = "pdfplumber"
except ImportError:
    try:
        import PyPDF2
        PDF_LIBRARY = "pypdf2"
    except ImportError:
        PDF_LIBRARY = None

# Try to import docx library
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def extract_text_from_pdf_pdfplumber(file_path: Union[str, Path]) -> str:
    """Extract text from PDF using pdfplumber."""
    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_from_pdf_pypdf2(file_path: Union[str, Path]) -> str:
    """Extract text from PDF using PyPDF2."""
    text_parts = []
    with open(file_path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_from_pdf_bytes_pdfplumber(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using pdfplumber."""
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_from_pdf_bytes_pypdf2(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using PyPDF2."""
    text_parts = []
    reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_from_docx(file_path: Union[str, Path]) -> str:
    """Extract text from DOCX file."""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx library not installed. Install with: pip install python-docx")

    doc = docx.Document(file_path)
    text_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)
    return "\n".join(text_parts)


def extract_text_from_docx_bytes(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes."""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx library not installed. Install with: pip install python-docx")

    doc = docx.Document(io.BytesIO(file_bytes))
    text_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)
    return "\n".join(text_parts)


def extract_text_from_txt(file_path: Union[str, Path]) -> str:
    """Extract text from TXT file."""
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


def parse_cv(file_path: Union[str, Path] = None, file_bytes: bytes = None, filename: str = None) -> str:
    """
    Parse CV from file path or bytes.

    Args:
        file_path: Path to the CV file
        file_bytes: Raw bytes of the CV file
        filename: Original filename (needed when using bytes to determine type)

    Returns:
        Extracted text from the CV
    """
    if file_path:
        file_path = Path(file_path)
        suffix = file_path.suffix.lower()

        if suffix == ".pdf":
            if PDF_LIBRARY == "pdfplumber":
                return extract_text_from_pdf_pdfplumber(file_path)
            elif PDF_LIBRARY == "pypdf2":
                return extract_text_from_pdf_pypdf2(file_path)
            else:
                raise ImportError("No PDF library available. Install pdfplumber or PyPDF2.")

        elif suffix == ".docx":
            return extract_text_from_docx(file_path)

        elif suffix == ".txt":
            return extract_text_from_txt(file_path)

        else:
            raise ValueError(f"Unsupported file format: {suffix}")

    elif file_bytes and filename:
        suffix = Path(filename).suffix.lower()

        if suffix == ".pdf":
            if PDF_LIBRARY == "pdfplumber":
                return extract_text_from_pdf_bytes_pdfplumber(file_bytes)
            elif PDF_LIBRARY == "pypdf2":
                return extract_text_from_pdf_bytes_pypdf2(file_bytes)
            else:
                raise ImportError("No PDF library available. Install pdfplumber or PyPDF2.")

        elif suffix == ".docx":
            return extract_text_from_docx_bytes(file_bytes)

        elif suffix == ".txt":
            return file_bytes.decode("utf-8")

        else:
            raise ValueError(f"Unsupported file format: {suffix}")

    else:
        raise ValueError("Either file_path or (file_bytes and filename) must be provided")


def extract_skills_from_cv(cv_text: str) -> list:
    """
    Extract potential skills from CV text for MCQ generation.
    This is a simple extraction - the LLM will do better skill identification.
    """
    # Common skill keywords to look for
    skill_indicators = [
        "skills", "technical skills", "technologies", "tools",
        "programming", "languages", "frameworks", "libraries"
    ]

    skills = set()
    lines = cv_text.lower().split("\n")

    in_skills_section = False
    for line in lines:
        line_lower = line.lower().strip()

        # Check if we're entering a skills section
        for indicator in skill_indicators:
            if indicator in line_lower and len(line_lower) < 50:
                in_skills_section = True
                break

        # If in skills section, extract comma/pipe separated items
        if in_skills_section:
            # Common separators in skill lists
            for sep in [",", "|", "•", "·", ";"]:
                if sep in line:
                    parts = line.split(sep)
                    for part in parts:
                        part = part.strip()
                        if part and len(part) < 30:
                            skills.add(part)

            # Check if we're leaving skills section (empty line or new section)
            if not line.strip() or line.startswith("#") or ":" in line:
                in_skills_section = False

    return list(skills)
